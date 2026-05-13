from io import BytesIO

import fitz
from docx import Document as DocxDocument

from src.config.settings import settings
from src.document.parser import parse_document_bytes


async def test_markdown_parser_preserves_headings_and_code_blocks():
    markdown = b"""# Airflow Guide

## Scheduler
The scheduler heartbeat is important.

```python
print("hello")
```
"""

    parsed = await parse_document_bytes(markdown, "airflow-guide.md", "text/markdown")

    assert parsed.parser_name == "markdown_parser"
    assert parsed.byte_size == len(markdown)
    assert parsed.metadata["normalizer"]["tool"] == "passthrough"
    assert any(
        block.block_type == "heading" and block.section_path == "Airflow Guide > Scheduler"
        for block in parsed.blocks
    )
    code_blocks = [block for block in parsed.blocks if block.block_type == "code"]
    assert code_blocks
    assert code_blocks[0].contains_code is True
    assert code_blocks[0].section_path == "Airflow Guide > Scheduler"


async def test_pdf_parser_falls_back_to_local_normalizer(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "normalizer_cache_dir", str(tmp_path / "normalized"))
    monkeypatch.setattr(settings, "mineru_api_key", "")
    monkeypatch.setattr(settings, "normalizer_enable_fallback", True)
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "CrashLoopBackOff troubleshooting")
    pdf_bytes = pdf.tobytes()
    pdf.close()

    parsed = await parse_document_bytes(pdf_bytes, "k8s.pdf", "application/pdf")

    assert parsed.parser_name == "markdown_parser"
    assert parsed.byte_size == len(pdf_bytes)
    assert parsed.metadata["normalizer"]["tool"] == "mineru_fallback_local"
    assert parsed.blocks
    paragraph = next(block for block in parsed.blocks if block.block_type == "paragraph")
    assert paragraph.section_path == "k8s"
    assert "CrashLoopBackOff" in paragraph.text


async def test_docx_parser_falls_back_to_local_normalizer(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "normalizer_cache_dir", str(tmp_path / "normalized"))
    monkeypatch.setattr(settings, "mineru_api_key", "")
    monkeypatch.setattr(settings, "normalizer_enable_fallback", True)
    docx = DocxDocument()
    docx.add_heading("Runbook", level=1)
    docx.add_paragraph("Check service status and logs.")
    buffer = BytesIO()
    docx.save(buffer)

    parsed = await parse_document_bytes(
        buffer.getvalue(),
        "runbook.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert parsed.parser_name == "markdown_parser"
    assert parsed.byte_size == len(buffer.getvalue())
    assert parsed.metadata["normalizer"]["tool"] == "mineru_fallback_local"
    paragraph = next(block for block in parsed.blocks if block.block_type == "paragraph")
    assert paragraph.section_path == "Runbook"
    assert "service status" in paragraph.text
