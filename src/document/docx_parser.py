"""DOCX parser using python-docx."""
from __future__ import annotations

import asyncio
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument

from src.document.parser import ParsedBlock, ParsedDocument

PARSER_NAME = "python_docx"
PARSER_VERSION = "1.0"


def _parse_docx_sync(file_bytes: bytes, file_name: str, content_type: str) -> ParsedDocument:
    doc = DocxDocument(BytesIO(file_bytes))
    blocks: list[ParsedBlock] = []
    section_stack: list[str] = []
    cursor = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.lower().startswith("heading"):
            try:
                level = int(style_name.split()[-1])
            except ValueError:
                level = 1
            section_stack = section_stack[: level - 1]
            section_stack.append(text)
            blocks.append(
                ParsedBlock(
                    text=text,
                    block_type="heading",
                    heading_level=level,
                    section_path=" > ".join(section_stack),
                )
            )
            continue

        start_char = cursor
        end_char = start_char + len(text)
        blocks.append(
            ParsedBlock(
                text=text,
                block_type="paragraph",
                section_path=" > ".join(section_stack) or None,
                start_char=start_char,
                end_char=end_char,
            )
        )
        cursor = end_char + 1

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        text = "\n".join(row for row in rows if row.strip())
        if text:
            blocks.append(
                ParsedBlock(
                    text=text,
                    block_type="table",
                    section_path=" > ".join(section_stack) or None,
                    metadata={"has_table": True},
                )
            )

    return ParsedDocument(
        title=Path(file_name).stem,
        file_name=file_name,
        content_type=content_type,
        parser_name=PARSER_NAME,
        parser_version=PARSER_VERSION,
        blocks=blocks,
        metadata={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
    )


async def parse_docx(
    file_bytes: bytes,
    file_name: str,
    content_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
) -> ParsedDocument:
    return await asyncio.to_thread(_parse_docx_sync, file_bytes, file_name, content_type)
